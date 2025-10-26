import os, io, re, glob
from typing import List, Tuple, Dict, Optional
import chardet

# TXT
def _read_txt(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read()
    enc = chardet.detect(raw).get("encoding") or "utf-8"
    return raw.decode(enc, errors="replace")

# PDF – 1) pypdf szybkie, 2) pdfplumber dokładniejsze układanie
def _read_pdf(path: str) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
        text = "\n".join(text_parts)
        if text.strip():
            return text
    except Exception:
        pass

    # fallback pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        text = "\n".join([p.extract_text() or "" for p in reader.pages])
        return text
    except Exception as e:
        return f"[PDF ERROR] {e}"

# DOCX
def _read_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    # tabele
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            texts.append(" | ".join(cells))
    return "\n".join(texts)

# EXCEL (xlsx/xls) + CSV
def _read_excel_like(path: str) -> str:
    import pandas as pd
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        df = pd.read_csv(path, dtype=str, na_filter=False)
        return df.to_csv(index=False)
    else:
        xl = pd.ExcelFile(path)
        out = []
        for sheet in xl.sheet_names:
            df = xl.parse(sheet, dtype=str).fillna("")
            out.append(f"\n=== SHEET: {sheet} ===\n" + df.to_csv(index=False))
        return "\n".join(out)

# OCR dla zeskanowanych PDF (opcjonalnie)
def _read_pdf_ocr(path: str) -> str:
    try:
        from pdf2image import convert_from_path
        import pytesseract
        pages = convert_from_path(path, dpi=200)
        ocr_texts = []
        for img in pages:
            ocr_texts.append(pytesseract.image_to_string(img))
        return "\n".join(ocr_texts)
    except Exception as e:
        return f"[PDF OCR ERROR] {e}"

# Normalizacja prostego tekstu
_WS = re.compile(r"[ \t\u00A0]+")
def _clean_text(t: str) -> str:
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    t = _WS.sub(" ", t)
    # zredukuj wielokrotne puste linie
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()

# Publiczna funkcja ładująca jeden plik
def load_file_text(path: str, prefer_ocr_pdf: bool=False) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext in (".txt",):
        raw = _read_txt(path)
    elif ext in (".pdf",):
        raw = _read_pdf(path)
        if prefer_ocr_pdf and (not raw or len(raw.strip()) < 10):
            raw = _read_pdf_ocr(path)
    elif ext in (".docx",):
        raw = _read_docx(path)
    elif ext in (".xlsx", ".xls", ".csv"):
        raw = _read_excel_like(path)
    else:
        return f"[UNSUPPORTED] Rozszerzenie {ext} nieobsługiwane."
    return _clean_text(raw or "")

# Prosty chunking
def chunk_text(text: str, max_chars: int = 3000, overlap: int = 200) -> List[str]:
    if max_chars <= 0:
        return [text]
    chunks = []
    n = len(text)
    i = 0
    while i < n:
        end = min(n, i + max_chars)
        chunk = text[i:end]
        chunks.append(chunk)
        if end == n:
            break
        i = end - overlap if overlap > 0 else end
    return chunks

# Ładowanie plików z katalogu / wzorca glob
def load_many(paths_or_globs: List[str], prefer_ocr_pdf: bool=False) -> Dict[str, str]:
    collected = {}
    for pattern in paths_or_globs:
        matches = []
        if os.path.isdir(pattern):
            for root, _, files in os.walk(pattern):
                for fn in files:
                    if os.path.splitext(fn)[1].lower() in (".txt",".pdf",".docx",".xlsx",".xls",".csv"):
                        matches.append(os.path.join(root, fn))
        else:
            matches = glob.glob(pattern)
        for p in matches:
            try:
                collected[p] = load_file_text(p, prefer_ocr_pdf=prefer_ocr_pdf)
            except Exception as e:
                collected[p] = f"[LOAD ERROR] {e}"
    return collected
